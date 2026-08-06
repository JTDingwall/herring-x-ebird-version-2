"""v31: apply the fifteen author comments on v30.

Longer sentences and paragraphs, ordinary connectives restored, no field names,
no pipeline vocabulary, no deferring cross-references, no restated conditions.
Comparator and control detail moves out of the Introduction. The conservation
sentence the author asked for is drafted rather than handed back. The
date-precision audit and the six period definitions become Table 1.

No number, citation or result changes.
"""
import re
import copy
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "mer_manuscript_v30.docx"
DST = "mer_manuscript_v31.docx"

d = docx.Document(SRC)
P = d.paragraphs


def set_text(i, s):
    p = P[i]
    for r in list(p.runs)[1:]:
        r._element.getparent().remove(r._element)
    if not p.runs:
        p.add_run("")
    p.runs[0].text = s


def italicise(i, phrases):
    p = P[i]
    text = "".join(r.text for r in p.runs)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    pat = "(" + "|".join(re.escape(x) for x in phrases) + ")"
    for piece in re.split(pat, text):
        if not piece:
            continue
        r = p.add_run(piece)
        r.italic = piece in phrases
        r.bold = False


# =====================================================================
# INTRODUCTION
# =====================================================================

# id=2: paragraphs 8 and 9 overlap. Combine into one.
set_text(8,
"Pacific herring (Clupea pallasii) spawning is one of the largest recurrent food "
"pulses on this coast. Adults crowd into shallow water in late winter and early "
"spring and attach adhesive eggs onto marine vegetation in the intertidal and "
"shallow subtidal (Haegele and Schweigert 1985; Dingwall et al. 2026). Spawning "
"at any one shoreline can pass quickly, and on the day of spawn the fish are "
"vulnerable and very close to shore, with some left stranded as the tide "
"recedes. The eggs then remain on marine vegetation for around two weeks before "
"hatching or washing away. Although the same general regions are used again from "
"year to year, the exact shoreline, extent and date shift among years (Hay and "
"Kronlund 1987; Hay et al. 2009; Rooper et al. 2024), so the pulse is "
"predictable at the scale of a coastal region but much less so at any particular "
"beach. What is available also changes as an event proceeds, because adult fish "
"create a short feeding opportunity when they enter the shallows, whereas "
"attached eggs persist after the adults have left, and eggs on floating "
"vegetation or exposed by a falling tide remain reachable by other routes.")
italicise(8, ["Clupea pallasii"])
set_text(9, "")

# id=3: remove the genus enumeration. Name the birds, not the genera.
set_text(10,
"Birds reach this food in different ways. Surf Scoter (Melanitta perspicillata), "
"White-winged Scoter (Melanitta deglandi) and Harlequin Duck (Histrionicus "
"histrionicus) dive for attached eggs, whereas loons, mergansers, grebes and "
"cormorants pursue fish and are therefore expected to concentrate on adults, "
"although they may take eggs as well. Gulls are more flexible again, taking eggs "
"from floating vegetation and herring pushed to the surface or into very shallow "
"water, as well as eggs and spent fish stranded along the shore. Bald Eagles and "
"corvids work the shoreline, shorebirds take eggs exposed by a falling tide, and "
"alcids pursue fish in open water and have little reason to gather at a "
"shoreline deposit at all.")
italicise(10, ["Melanitta perspicillata", "Melanitta deglandi",
               "Histrionicus histrionicus"])

# id=7: comparator and control detail belongs in Methods, where it already is.
set_text(14,
"This study links the two records for 49 coastal bird species. It asks whether "
"the difference between checklists within 5 km of a recorded spawning event and "
"those 5 to 20 km away widened after the event date, compared with the preceding "
"two weeks. Whether a species appeared on a checklist and how many birds were "
"counted when an observer supplied a number are examined separately, because a "
"species can become more widespread among checklists without forming larger "
"recorded groups, or form larger groups without appearing on more checklists.")

# id=9: the author asked for this sentence, so write it.
set_text(15,
"Species with established or plausible herring links were expected to be "
"reported more often, counted in larger numbers, or both once spawning began, "
"and the feeding groups were expected to differ in when they responded. Gulls "
"and shoreline scavengers should respond near the event date, whereas diving "
"ducks that feed on attached roe should respond once eggs are available, and "
"alcids that feed away from the shoreline deposit were expected to show little "
"response at all. Testing all 49 species places the well-studied birds inside a "
"broader coastal community and identifies which species, and which time windows, "
"would repay focused fieldwork. That matters beyond the birds themselves. "
"Herring on this coast are managed from spawn surveys that were never designed "
"to describe the predators a deposit supports, so a regional account of which "
"species track spawning, and when, would give fisheries management and coastal "
"bird conservation a shared description of what a spawning event is worth.")

# =====================================================================
# METHODS
# =====================================================================

# id=11: nobody cares about study rules.
set_text(23,
"Records came from the eBird Basic Dataset released in May 2026 and covered 2005 "
"to 2025 (eBird 2026). Eligible checklists used stationary or travelling "
"protocols, lasted 5 to 300 minutes, covered no more than 5 km, and had one to "
"ten observers. Only complete checklists were retained, because marking a "
"checklist complete means the observer reported every species they detected and "
"identified, so a species that does not appear on one can be treated as a "
"genuine non-detection rather than as missing information.")

# id=12 and id=18: describe how the data are collected; no field names.
set_text(25,
"Herring records came from the Pacific Herring Spawn Index, compiled by "
"Fisheries and Oceans Canada (Fisheries and Oceans Canada 2026; Grinnell et al. "
"2023). Surveyors work the coast through the spawning season and record where "
"spawn has been deposited, using surface surveys that measure the length of "
"shoreline spawned and the extent of eggs on vegetation, and dive surveys that "
"additionally record egg layering and depth. Each record therefore carries a "
"location together with the dates over which spawn was observed at that place. "
"The resulting index is relative rather than absolute, because it describes what "
"surveyors found, and coastline that was not surveyed yields no record rather "
"than a zero. Only the presence, location and timing of a recorded event entered "
"the models, and the index value itself was not used as a measure of how much "
"spawn was deposited.")

# id=16: combine 27 and 28, and drop the phrasing flagged as unscientific.
set_text(27,
"Linking the two sources produced 217,200 complete eBird checklists and 1,120 "
"recorded spawning events, which fell into 58 event blocks, 29,248 observer "
"clusters and 22,980 generalized location clusters. An event block groups "
"recorded spawning events that are connected because they link to the same "
"checklists at the same time, which keeps those events and the checklists that "
"reach them under a single random effect. Blocks are a statistical grouping and "
"carry no claim about the biology, since two events in the same block need not "
"belong to one spawning population. A checklist linked to several events "
"remained one row in the model and carried the random effect for its block.")
set_text(28, "")

# id=17 and id=18: unreadable. Short paragraph, definitions move to Table 1.
set_text(29,
"Day 0 was the midpoint between the first and last dates on which spawn was "
"recorded at a location, rounded down to a whole day, or that single date where "
"only one was available. It is an event-date anchor and not an observed "
"biological onset, and it is dated more precisely for some events than for "
"others: among the 1,144 events in the wider linked frame, 466 had spawn "
"recorded on a single day and the median span between first and last date was "
"one day. Six periods were defined around that anchor, and Table 1 gives their "
"limits, their durations and the food expected to be available in each.")

# =====================================================================
# 2.3  four short paragraphs become two
# =====================================================================
set_text(37,
"Two outcomes represent different parts of the checklist record. Checklist "
"reporting was one when a species was reported on an eligible complete checklist "
"and zero when it was omitted under the fixed ambiguity rules, and although an "
"omission is informative it does not prove that the species was absent. Every "
"positive report kept its count state, so a finite number contributed to both "
"outcomes while an unquantified X contributed only to whether the species was "
"reported. Observers enter X when they know a bird is present but cannot or will "
"not count it, most often when the flock is large.")

# id=26: the restated condition is deleted.
set_text(38,
"The second outcome was the natural logarithm of the number reported. "
"Exponentiated estimates therefore compare geometric mean counts among those "
"records, and they do not measure abundance across all checklists, since a total "
"from a travelling checklist need not represent a single flock. The two outcomes "
"answer different questions. The first asks whether a species appeared on a "
"checklist at all, whereas the second asks whether the number an observer "
"supplied was larger, and a species can change on one without changing on the "
"other. Because the count outcome is measured only among birds that were both "
"reported and counted, however, it is conditioned on something the reporting "
"models themselves show to be responsive, which is a selection problem and not "
"only a difference of question.")
set_text(39, "")
set_text(40, "")

# id=29: the deferring cross-reference is deleted.
set_text(49,
"Whether a species was reported was fitted with a binomial mixed model, and log "
"count with a Gaussian mixed model, both using the nAGQ = 0 approximation for "
"the binomial fits (Bates et al. 2015). Confidence intervals are Wald intervals "
"on the model scale, built from the exact variance of each comparison against "
"the full fixed-effect covariance matrix and then exponentiated, so both the "
"estimates and their intervals inherit that approximation.")

# id=31: no pipeline vocabulary.
set_text(50,
"Models adjusted for year as a factor, protocol, log duration, log of one plus "
"distance travelled, and number of observers, and carried random intercepts for "
"event block, observer and location. Day of year and time of day were not "
"available in the linked dataset, so seasonal balance between the zones could "
"not be demonstrated and time of day could not be adjusted for; the absence of a "
"diel term is a real gap rather than a considered omission. The three effort "
"terms carry their own caution, because duration, distance and observer count "
"are measured on the same checklist as the outcome, and a conspicuous flock may "
"itself lengthen a visit, extend a route or draw a second observer, in which "
"case adjusting for them removes part of the response. Their distributions "
"barely move across the event date, however: median near-zone duration rises by "
"one minute and median near-zone distance by 0.04 km between the two windows, "
"with slightly larger shifts in the reference zone. That is reassuring without "
"being decisive, since a covariate can respond to exposure without its marginal "
"distribution shifting.")

# id=8: American Robin now has its first mention here, so it needs a binomial.
set_text(57, P[57].text.replace(
    "with American Robin as a terrestrial control",
    "with American Robin (Turdus migratorius) as a terrestrial control"))
italicise(57, ["Turdus migratorius"])

d.save(DST)

# =====================================================================
# Table 1, inserted after the Day 0 paragraph
# =====================================================================
d = docx.Document(DST)
P = d.paragraphs
anchor = [k for k, p in enumerate(P)
          if p.text.startswith("Day 0 was the midpoint")][0]
after = P[anchor]

ROWS = [
    ("Period", "Days from anchor", "Days", "Expected food"),
    ("Baseline", "−28 to −15", "14", "No spawn recorded"),
    ("Early pre-spawn", "−14 to −8", "7", "No spawn recorded"),
    ("Immediate pre-spawn", "−7 to −1", "7", "No spawn recorded"),
    ("Spawn", "0 to 3", "4", "Adults in the shallows; eggs being laid"),
    ("Early egg", "4 to 14", "11", "Eggs on vegetation and along the shore"),
    ("Late egg", "15 to 28", "14", "Eggs largely hatched or washed away"),
]

t = d.add_table(rows=0, cols=4)
t.style = "Table Grid"
for j, row in enumerate(ROWS):
    cells = t.add_row().cells
    for k, val in enumerate(row):
        cells[k].text = ""
        r = cells[k].paragraphs[0].add_run(val)
        r.bold = (j == 0)
        r.font.size = Pt(9)

cap = d.add_paragraph()
cr = cap.add_run(
    "Table 1. Periods defined relative to the recorded event date. The pre-spawn "
    "summary weights the two pre-spawn periods equally. The active summary "
    "weights the spawn and early egg periods by their durations, 4/15 and 11/15.")
cr.font.size = Pt(9)

after._element.addnext(cap._element)
after._element.addnext(t._element)

d.save(DST)

# =====================================================================
# drop the emptied paragraphs, protecting images and equations
# =====================================================================
d = docx.Document(DST)
NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
NS_M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
removed = 0
for p in list(d.paragraphs):
    if p.text.strip() or p.style.name != "Normal":
        continue
    if (p._element.findall(".//" + NS_A)
            or p._element.findall(".//" + NS_M + "oMath")
            or p._element.findall(".//" + NS_M + "oMathPara")):
        continue
    p._element.getparent().remove(p._element)
    removed += 1
d.save(DST)
print("wrote", DST, "| empty paragraphs removed:", removed)

# =====================================================================
# Second pass: the clippiness is paper-wide, not only where the author
# commented. Results 3.1 to 3.4 average 16.6 to 20.0 words a sentence with
# up to a third under twelve. Merge with connectives. No number changes.
# =====================================================================
d = docx.Document(DST)
P = d.paragraphs
idx = {}
for k, p in enumerate(P):
    t = p.text.strip()
    if t:
        idx[t[:60]] = k


def rewrite(prefix, new):
    for k, p in enumerate(P):
        if p.text.strip().startswith(prefix):
            for r in list(p.runs)[1:]:
                r._element.getparent().remove(r._element)
            if not p.runs:
                p.add_run("")
            p.runs[0].text = new
            return k
    raise SystemExit("not found: " + prefix)


rewrite("The analysis covered 217,200",
"The analysis covered 217,200 complete eBird checklists collected between 2005 "
"and 2025, linked to 1,120 recorded spawning events that formed 58 blocks. How "
"much evidence stands behind each species varies enormously, since American Crow "
"(Corvus brachyrhynchos) appeared on 53.5% of checklists and carried 112,180 "
"counted records, whereas Glaucous Gull (Larus hyperboreus) appeared on 0.09% "
"and carried 185.")

rewrite("Birds were counted in larger numbers far more",
"Birds were counted in larger numbers far more consistently than they were "
"reported more often (Figure 2). Reporting rose for 28 of 48 species and fell "
"for 20, of which 13 increases survived correction for multiple testing, whereas "
"counts rose for 42 of 46 species and fell for four, with 18 increases "
"surviving, and no decrease survived in either outcome. The gap between the two "
"outcomes is wider than those tallies suggest, because among the 46 species with "
"both models 27 rose in both, 15 rose only in counts, four rose in neither, and "
"none rose only in reporting, so every species that disagreed with itself "
"disagreed in the same direction. An exact paired test on those 15 gives "
"p = 0.00006, although species share checklists, events and observers, so the "
"figure describes how lopsided the pattern is more than it tests it.")

rewrite("Gulls dominated the reporting results",
"Gulls dominated the reporting results, with Bonaparte's Gull (Chroicocephalus "
"philadelphia) reported 1.48 times as often per nearby spawn (95% CI "
"1.22–1.79), American Herring Gull (L. smithsonianus) 1.39 (1.23–1.57), "
"California Gull (L. californicus) 1.32 (1.21–1.43), Iceland Gull "
"(L. glaucoides) 1.29 (1.18–1.41) and Short-billed Gull (L. brachyrhynchus) "
"1.25 (1.17–1.34). These are not 48% or 39% increases in the number of gulls on "
"the coast, since each describes how much more strongly a checklist's reporting "
"of that species tracked a nearby spawning event after the event date than "
"before it. Outside the gulls the largest reporting increases came from Lesser "
"Scaup (Aythya affinis) at 1.26 (1.11–1.44) and Northern Pintail (Anas acuta) "
"at 1.21 (1.10–1.32), neither of which is a bird the herring literature would "
"nominate first.")

rewrite("The largest count increases were Long-tailed Duck",
"The largest count increases were Long-tailed Duck (Clangula hyemalis) at 1.43 "
"(1.29–1.60), Surf Scoter at 1.30 (1.21–1.40), Short-billed Gull at 1.30 "
"(1.23–1.37), Harlequin Duck at 1.20 (1.15–1.26) and Common Goldeneye "
"(Bucephala clangula) at 1.20 (1.15–1.25), although increases were far more "
"widespread than that significant subset conveys. Of 46 species, only Common "
"Murre (Uria aalge), Black-bellied Plover (Pluvialis squatarola), Black "
"Oystercatcher (Haematopus bachmani) and Great Blue Heron (Ardea herodias) were "
"counted in smaller numbers at all, and none of those four significantly. The "
"tightest estimates are not the largest ones, because Bald Eagle (Haliaeetus "
"leucocephalus) shows a count increase of 1.063 (1.045–1.080) that is modest in "
"size and among the most precise in the family, resting on 83,359 counted "
"records, whereas Bonaparte's Gull's count increase of 1.16 spans 0.87 to 1.56 "
"on 4,025.")

rewrite("The 49 species were sorted into seven feeding groups", "")

rewrite("Roe-feeding diving sea ducks, the group with the strongest",
"The 49 species were sorted into seven feeding groups in the species registry "
"before any model was fitted, and those groups diverged sharply, with two "
"producing nothing at all. Roe-feeding diving sea ducks, the group with the "
"strongest prior expectation, gave the clearest count result, since seven of ten "
"were counted in significantly larger numbers: Long-tailed Duck 1.43 "
"(1.29–1.60), Surf Scoter 1.30 (1.21–1.40), Harlequin Duck 1.20 (1.15–1.26), "
"Common Goldeneye 1.20 (1.15–1.25), White-winged Scoter 1.17 (1.06–1.30), "
"Greater Scaup (A. marila) 1.16 (1.05–1.29) and Bufflehead (B. albeola) 1.12 "
"(1.09–1.15), while Barrow's Goldeneye (B. islandica) at 1.10 (1.01–1.19) fell "
"just outside and only Black Scoter (Melanitta americana) was flat. Reporting "
"behaved quite differently in the same birds, because only Lesser Scaup 1.26 "
"(1.11–1.44), Harlequin Duck 1.15 (1.05–1.26) and Greater Scaup 1.13 "
"(1.03–1.24) cleared correction, and Surf Scoter reporting sat at 1.00. Birds "
"that already winter on these shorelines and then concentrate where roe is "
"reachable would produce exactly this shape.")

rewrite("Gulls were the mirror image",
"Gulls were the mirror image, and the only group in which reporting was the "
"stronger outcome. Six of eight were reported significantly more often, from "
"Bonaparte's Gull at 1.48 down to Glaucous-winged Gull (L. glaucescens) at 1.14 "
"(1.08–1.20), with only Ring-billed and Western Gull outside, and four were also "
"counted in larger numbers, led by Short-billed Gull at 1.30 and "
"Glaucous-winged Gull at 1.17 (1.13–1.21). Bonaparte's Gull, which had the "
"largest reporting increase in the study, nonetheless had a count estimate of "
"1.16 spanning 0.87 to 1.56, so its numbers did not detectably change. Gulls are "
"also the group where an observation effect is hardest to rule out, because a "
"feeding flock is visible from a long way off and easy to record without "
"counting.")

rewrite("Fish-chasing piscivores showed a narrower version",
"Fish-chasing piscivores showed a narrower version of the sea duck pattern, with "
"five of twelve counted in larger numbers: Pacific Loon (Gavia pacifica) 1.16 "
"(1.04–1.29), Red-breasted Merganser (Mergus serrator) 1.13 (1.07–1.19), Common "
"Merganser (M. merganser) 1.11 (1.07–1.16), Double-crested Cormorant "
"(Nannopterum auritum) 1.10 (1.05–1.15) and Common Loon (G. immer) 1.05 "
"(1.01–1.08). Only Common Merganser at 1.16 (1.10–1.23) was also reported more "
"often, and seven of the twelve, including all three grebes, showed nothing at "
"all. Shoreline scavengers were narrower still, because Bald Eagle's count "
"increase of 1.06 (1.05–1.08) was the group's only significant result, while "
"American Crow, Common Raven (C. corax) and Great Blue Heron stayed flat in both "
"outcomes despite being three of the four most frequently reported birds in the "
"family. Since their estimates are precise, those flat results carry "
"information.")

rewrite("Two groups produced nothing, and one of them had a reason",
"Two groups produced nothing, and one of them had a reason to respond. No alcid "
"cleared correction in either outcome, since Pigeon Guillemot (Cepphus columba), "
"Marbled Murrelet (Brachyramphus marmoratus), Common Murre and Rhinoceros Auklet "
"had reporting estimates of 0.99, 0.94, 0.87 and 0.75, three of the four below "
"one, with intervals too wide to separate from no change. Birds that chase fish "
"in open water should not gather at a shoreline deposit, so that is the expected "
"answer. Shorebirds are the harder case, because they are documented consumers "
"of herring eggs on exposed substrate (Bishop and Green 2001), yet Black "
"Oystercatcher, Dunlin (C. alpina), Black Turnstone (Arenaria melanocephala) and "
"Black-bellied Plover were all reported slightly less often, between 0.94 and "
"0.96, none significantly, and none was counted in larger numbers either. "
"Surfbird had the group's largest reporting estimate at 1.37 but an interval "
"from 0.97 to 1.94 and no count model. Egg availability to a wader depends on "
"the tide, and these models carry no tidal or diel term, so a response confined "
"to the hours around low water would wash out across every other checklist, "
"while shorebird use of spawn here may also be too local to register across a "
"whole strait. The models fitted and simply found nothing.")

rewrite("Dabbling ducks and geese responded mainly through reporting",
"Dabbling ducks and geese responded mainly through reporting, with Northern "
"Pintail 1.21 (1.10–1.32), Brant (B. bernicla) 1.16 (1.04–1.30) and American "
"Wigeon (Mareca americana) 1.11 (1.04–1.17) clearing correction, as did Mallard "
"(A. platyrhynchos)'s count increase of 1.07 (1.04–1.10), while Canada Goose "
"(Branta canadensis) stayed flat. This is the group in which a herring-specific "
"reading of the family is hardest to defend.")

rewrite("The active-period result does not describe one common pattern",
"The active-period result does not describe one common pattern through time "
"(Figure 3), and the difference between groups can be tested. For each species, "
"the estimate at the event date minus the estimate during early egg gives a "
"timing score, whose variance comes from the fitted covariance matrix, and "
"regressing that score on the feeding groups fixed before any model was run, "
"weighted by precision, tests whether the groups differ. Both outcomes say they "
"do, counts far more decisively than reporting (Q = 117.3 and Q = 15.0, each on "
"6 degrees of freedom, p < 0.001 and p = 0.02), although the two outcomes "
"separate different groups. For counts the result comes from roe-feeding divers, "
"which peak once eggs are available, and from shoreline scavengers, which peak "
"at the event date, whereas for reporting the two groups whose intervals exclude "
"zero are gulls and shoreline scavengers, both peaking early, while the divers "
"that dominate the count result are flat. Feeding group therefore explains part "
"of the variation and not most of it, leaving 76% of the residual heterogeneity "
"unexplained for counts and 39% for reporting.")

rewrite("Individual species show the same division",
"Individual species show the same division. Common Merganser rose in both "
"outcomes with both peaks close to the event date, Glaucous-winged and "
"Short-billed Gull rose in both with the reporting change concentrated early, "
"and Bald Eagle rose briefly near the event date in both. Surf Scoter reporting "
"stayed at one throughout while its counts peaked during early egg availability, "
"and White-winged Scoter and Harlequin Duck traced the same shape with smaller "
"amplitude. Figure 3 shows these eight, chosen to span the feeding groups.")

rewrite("American Wigeon appears in Figure 3 as a counterpoint",
"American Wigeon appears in Figure 3 as a counterpoint, since its estimates were "
"modestly positive after the event date even though no direct herring link is "
"established for it, and its profile lacks the late peak that the roe-feeding "
"divers show.")

rewrite("Coastal birds were recorded differently around herring",
"Coastal birds were recorded differently around herring spawning events, and the "
"clearest signal was in how many birds observers wrote down rather than in how "
"often they wrote anything at all. That distinction is the main result, and it "
"holds up in three ways: it survived every change in how overlapping events were "
"represented, it is the outcome for which the timing test is most decisive, and "
"it is the harder of the two to produce through observer behaviour alone, "
"because writing a larger number means seeing more birds in a way that writing a "
"species name does not.")

rewrite("The timing adds detail a single before-and-after",
"The timing adds detail that a single before-and-after number cannot. Diving sea "
"ducks reached their largest counts once eggs were available, whereas shoreline "
"scavengers peaked as spawning began in both outcomes, and gulls peaked early in "
"reporting while showing no timing difference in counts. A test on feeding "
"groups fixed before fitting rejects equal timing in both outcomes, and the "
"groups it separates differ between them. This is what a spawn should look like "
"if it is not one uniform meal, because adult fish, attached eggs, dislodged "
"eggs, carcasses and the scavengers those attract become available to different "
"birds at different times and depths, and because a gull flock is conspicuous "
"enough to change whether a checklist mentions the species while a raft of "
"scoters changes how many get counted. The count evidence is much the stronger "
"of the two, however, and the reporting result should be read as its weaker "
"companion.")

rewrite("The analysis remains observational",
"The analysis remains observational. It accounts for seasonal change shared by "
"near and far shorelines and for their average difference before the event date, "
"but it cannot remove a change that differs between zones and happens to "
"coincide with spawning, and it never observes a bird eating herring, moving "
"toward a spawn, or being present at all. What it observes is what people wrote "
"on checklists.")

rewrite("Reporting and counting sit at different stages",
"Reporting and counting sit at different stages of the same process, since "
"reporting needs a bird to be present, available, detected, identified and "
"written down, whereas the count model starts after all of that and only when "
"the observer supplied a number. A species can therefore turn up in larger "
"flocks without appearing on more checklists, as Surf Scoter did, or the "
"reverse. That ordering matters for the outcome carrying the main conclusion, "
"because counts are measured among records that were made and quantified, and "
"reporting itself rose for 13 species, so for those species the checklists "
"contributing counts after the event date are not the set that would have "
"contributed before. Part of the count increase may therefore describe which "
"checklists carried a number rather than how many birds were on them. The "
"concern is weakest where reporting did not move, and Surf Scoter, flat in "
"reporting and up 1.30 in counts, is the clearest such case, while for species "
"in which both outcomes moved the direction of the bias is unknown.")

rewrite("The nearest-event sensitivity changed more reporting",
"The nearest-event sensitivity changed more reporting conclusions than count "
"conclusions. One reading is that flock size at a heavily used site is a more "
"concentrated signal than the line between reporting and not reporting across a "
"varied coastline, while another is that reporting odds are simply more "
"sensitive to how overlapping event footprints are assigned. The sensitivity "
"cannot say which dominates, although it does mark where species-level claims "
"need restraint.")

rewrite("An entry of X adds a third stage",
"An entry of X adds a third stage, because observers may write X when a flock is "
"large, distant, mobile, mixed or simply not counted. There was no family-wide "
"shift in that habit around spawning, which is mildly reassuring, but 30 of 41 "
"of those fits were singular and the check has little power. It would be wrong "
"to conclude that excluding X records makes the count estimates conservative, "
"since if observers write X more often when a flock is large then the counted "
"subset systematically misses the biggest groups, and misses them most often "
"when the birds are most concentrated, which would pull the count estimates "
"down. If X instead marks distant or hard-to-identify birds, the omission has no "
"particular relationship to flock size, and these data cannot separate the two "
"cases.")

d.save(DST)

d = docx.Document(DST)
for p in list(d.paragraphs):
    if p.text.strip() or p.style.name != "Normal":
        continue
    if (p._element.findall(".//" + NS_A)
            or p._element.findall(".//" + NS_M + "oMath")
            or p._element.findall(".//" + NS_M + "oMathPara")):
        continue
    p._element.getparent().remove(p._element)
d.save(DST)
print("second pass done")

# =====================================================================
# The rewrite helper collapses runs, so every scientific name in a
# rewritten paragraph lost its italics. Re-italicise globally.
# =====================================================================
import csv as _csv
REG = {}
with open("/sessions/keen-confident-allen/mnt/herring-x-ebird-version-2/"
          "metadata/canonical_species_registry.csv", encoding="utf-8-sig") as fh:
    for r in _csv.DictReader(fh):
        REG[r["common_name"]] = r["scientific_name"]

ALL = ["Clupea pallasii", "Turdus migratorius"]
for sci in REG.values():
    g, s = sci.split()[0], sci.split()[1]
    ALL.append(sci)
    if f"{g[0]}. {s}" != "M. americana":
        ALL.append(f"{g[0]}. {s}")
ALL = sorted(set(ALL), key=len, reverse=True)
PAT = "(" + "|".join(re.escape(x) for x in ALL) + ")"

d = docx.Document(DST)
P = d.paragraphs
ri = [k for k, p in enumerate(P) if p.text.strip() == "References"][0]
n_ital = 0
for k in range(5, ri):                      # skip title block and abstract
    p = P[k]
    text = "".join(r.text for r in p.runs)
    if not re.search(PAT, text):
        continue
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    for piece in re.split(PAT, text):
        if not piece:
            continue
        r = p.add_run(piece)
        r.italic = bool(re.fullmatch(PAT, piece))
        r.bold = False
        if r.italic:
            n_ital += 1
d.save(DST)
print("scientific names re-italicised:", n_ital)

# =====================================================================
# Moving the comparator detail out of the Introduction moved four first
# mentions. Put the binomials where the species now first appears and
# remove the copies that have become second mentions.
# =====================================================================
d = docx.Document(DST)


def sub(prefix, old, new):
    for p in d.paragraphs:
        t = "".join(r.text for r in p.runs)
        if t.strip().startswith(prefix) and old in t:
            t = t.replace(old, new, 1)
            for r in list(p.runs)[1:]:
                r._element.getparent().remove(r._element)
            p.runs[0].text = t
            return True
    raise SystemExit("sub failed: " + prefix + " | " + old)


sub("Birds reach this food", "Bald Eagles and corvids",
    "Bald Eagles (Haliaeetus leucocephalus) and corvids")
sub("Gadwall and Northern Shoveler were carried",
    "Gadwall and Northern Shoveler",
    "Gadwall (Mareca strepera) and Northern Shoveler (Spatula clypeata)")
sub("To test whether the 5 km near-zone boundary",
    "Glaucous-winged Gull were examined",
    "Glaucous-winged Gull (Larus glaucescens) were examined")
sub("The analysis covered 217,200", "Glaucous Gull (Larus hyperboreus)",
    "Glaucous Gull (L. hyperboreus)")
sub("The largest count increases were Long-tailed Duck",
    "Bald Eagle (Haliaeetus leucocephalus) shows", "Bald Eagle shows")
sub("Gulls were the mirror image",
    "Glaucous-winged Gull (L. glaucescens) at 1.14", "Glaucous-winged Gull at 1.14")
d.save(DST)

d = docx.Document(DST)
P = d.paragraphs
ri = [k for k, p in enumerate(P) if p.text.strip() == "References"][0]
n2 = 0
for k in range(5, ri):
    p = P[k]
    text = "".join(r.text for r in p.runs)
    if not re.search(PAT, text):
        continue
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    for piece in re.split(PAT, text):
        if not piece:
            continue
        r = p.add_run(piece)
        r.italic = bool(re.fullmatch(PAT, piece))
        r.bold = False
        if r.italic:
            n2 += 1
d.save(DST)
print("binomial relocation done; italics:", n2)
